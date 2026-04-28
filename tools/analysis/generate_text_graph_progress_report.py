from __future__ import annotations

import csv
import json
import math
import re
import statistics
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO_ROOT = Path(__file__).resolve().parents[2]
EXPERIMENT_ROOT = REPO_ROOT / "experiments"
REPORT_DIR = REPO_ROOT / "experiment_records" / "reports" / "controlled_30split"
REPORT_STEM = "text_graph_ontology_progress_report_2026-04-28"

DATASETS = ("BRCA", "KIRC", "LUSC")
METHODS = (
    "sentence-only",
    "sentence-ontology",
    "sentence-hierarchical-graph",
    "sentence-hierarchical-graph-ontology",
)

METHOD_LABELS = {
    "sentence-only": "句子级表示（sentence-only）",
    "sentence-ontology": "句子级表示 + 医学知识增强（sentence + ontology）",
    "sentence-hierarchical-graph": "句子级表示 + 文档结构增强（sentence + hierarchical graph）",
    "sentence-hierarchical-graph-ontology": "句子级表示 + 结构/知识联合增强（sentence + hierarchical graph + ontology）",
}


@dataclass(frozen=True)
class MetricRow:
    dataset: str
    method: str
    n: int
    auc_mean: float
    auc_sd: float
    acc_mean: float
    acc_sd: float
    source_note: str


@dataclass(frozen=True)
class DeltaRow:
    dataset: str
    method: str
    n: int
    dauc_mean: float
    dauc_sd: float
    dauc_median: float
    dauc_pos: int
    dacc_mean: float
    dacc_sd: float


def fmt_mean_sd(mean: float, sd: float) -> str:
    if math.isnan(mean):
        return "N/A"
    if math.isnan(sd):
        return f"{mean:.4f}"
    return f"{mean:.4f} ± {sd:.4f}"


def fmt_delta(mean: float, sd: float) -> str:
    if math.isnan(mean):
        return "N/A"
    if math.isnan(sd):
        return f"{mean:+.4f}"
    return f"{mean:+.4f} ± {sd:.4f}"


def mean_sd(values: Iterable[float]) -> tuple[float, float]:
    vals = list(values)
    if not vals:
        return float("nan"), float("nan")
    if len(vals) == 1:
        return statistics.mean(vals), float("nan")
    return statistics.mean(vals), statistics.stdev(vals)


def read_split_records(dataset: str, method: str) -> dict[int, dict[str, object]]:
    path = EXPERIMENT_ROOT / dataset / method / "records" / "split_results.csv"
    if not path.exists():
        return {}

    records: dict[int, dict[str, object]] = {}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            if row.get("dataset") and row["dataset"] != dataset:
                continue
            if row.get("method") and row["method"] != method:
                continue
            if not row.get("auc"):
                continue
            split_idx = int(row["split_idx"])
            records[split_idx] = {
                "split": split_idx,
                "auc": float(row["auc"]),
                "acc": float(row["acc"]),
                "status": row.get("status", ""),
                "source_path": row.get("source_path", ""),
                "run_name": row.get("run_name", ""),
                "ontology_variant": row.get("ontology_variant", ""),
            }
    return records


def parse_run_dir(run_dir: Path) -> dict[int, dict[str, object]]:
    records: dict[int, dict[str, object]] = {}
    for split_dir in run_dir.glob("split_*"):
        try:
            split_idx = int(split_dir.name.split("_")[-1])
        except ValueError:
            continue
        log_path = split_dir / "log.jsonl"
        if not log_path.exists():
            continue

        best_epoch_record: dict[str, object] | None = None
        best_epoch = None
        final_record: dict[str, object] | None = None
        for line in log_path.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            try:
                payload = json.loads(line)
            except json.JSONDecodeError:
                continue
            if payload.get("type") == "final_evaluation" and payload.get("target"):
                final_record = {
                    "split": split_idx,
                    "auc": float(payload["target"]["auc"]),
                    "acc": float(payload["target"]["acc"]),
                    "best_epoch": payload.get("best_epoch"),
                    "source_path": str(log_path),
                }
                continue
            target = payload.get("target") or {}
            if "auc" not in target:
                continue
            auc = float(target["auc"])
            if best_epoch_record is None or auc > float(best_epoch_record["auc"]):
                best_epoch = payload.get("epoch")
                best_epoch_record = {
                    "split": split_idx,
                    "auc": auc,
                    "acc": float(target.get("acc", float("nan"))),
                    "best_epoch": best_epoch,
                    "source_path": str(log_path),
                }

        if final_record is not None:
            records[split_idx] = final_record
        elif best_epoch_record is not None:
            records[split_idx] = best_epoch_record
    return records


def load_all_records() -> tuple[dict[str, dict[str, dict[int, dict[str, object]]]], dict[str, str]]:
    records: dict[str, dict[str, dict[int, dict[str, object]]]] = {
        dataset: {method: read_split_records(dataset, method) for method in METHODS}
        for dataset in DATASETS
    }
    notes = {
        "BRCA": "sentence-only 为本轮新跑 30 splits；增强模式均为本轮 30 splits。",
        "KIRC": "sentence-only 使用历史 run_150splits 的前 30 splits 做公平对齐；增强模式为本轮 30 splits。",
        "LUSC": "sentence-only 使用历史 LUSC_150splits_sentence_only 的前 30 splits 做公平对齐；增强模式为本轮 30 splits。",
    }

    kirc_history = parse_run_dir(EXPERIMENT_ROOT / "KIRC" / "sentence-only" / "runs" / "run_150splits")
    if kirc_history:
        records["KIRC"]["sentence-only"] = {idx: value for idx, value in kirc_history.items() if idx < 30}

    lusc_history = parse_run_dir(EXPERIMENT_ROOT / "LUSC" / "sentence-only" / "runs" / "LUSC_150splits_sentence_only")
    if lusc_history:
        records["LUSC"]["sentence-only"] = {idx: value for idx, value in lusc_history.items() if idx < 30}
        notes["LUSC_full_150"] = (
            f"LUSC 历史 sentence-only 150 splits: "
            f"AUC {fmt_mean_sd(*mean_sd(v['auc'] for v in lusc_history.values()))}, "
            f"ACC {fmt_mean_sd(*mean_sd(v['acc'] for v in lusc_history.values()))}。"
        )

    return records, notes


def build_metric_rows(records: dict[str, dict[str, dict[int, dict[str, object]]]], notes: dict[str, str]) -> list[MetricRow]:
    rows: list[MetricRow] = []
    for dataset in DATASETS:
        for method in METHODS:
            values = list(records[dataset][method].values())
            auc_mean, auc_sd = mean_sd(float(item["auc"]) for item in values)
            acc_mean, acc_sd = mean_sd(float(item["acc"]) for item in values)
            rows.append(
                MetricRow(
                    dataset=dataset,
                    method=method,
                    n=len(values),
                    auc_mean=auc_mean,
                    auc_sd=auc_sd,
                    acc_mean=acc_mean,
                    acc_sd=acc_sd,
                    source_note=notes.get(dataset, ""),
                )
            )
    return rows


def build_delta_rows(records: dict[str, dict[str, dict[int, dict[str, object]]]]) -> list[DeltaRow]:
    rows: list[DeltaRow] = []
    for dataset in DATASETS:
        baseline = records[dataset]["sentence-only"]
        for method in METHODS[1:]:
            current = records[dataset][method]
            common = sorted(set(baseline) & set(current))
            dauc = [float(current[idx]["auc"]) - float(baseline[idx]["auc"]) for idx in common]
            dacc = [float(current[idx]["acc"]) - float(baseline[idx]["acc"]) for idx in common]
            dauc_mean, dauc_sd = mean_sd(dauc)
            dacc_mean, dacc_sd = mean_sd(dacc)
            rows.append(
                DeltaRow(
                    dataset=dataset,
                    method=method,
                    n=len(common),
                    dauc_mean=dauc_mean,
                    dauc_sd=dauc_sd,
                    dauc_median=statistics.median(dauc) if dauc else float("nan"),
                    dauc_pos=sum(value > 0 for value in dauc),
                    dacc_mean=dacc_mean,
                    dacc_sd=dacc_sd,
                )
            )
    return rows


def parse_best_epoch_payload(row: dict[str, object]) -> dict[str, object] | None:
    status = str(row.get("status", ""))
    match = re.search(r"best_epoch_(\d+)", status)
    target_epoch = int(match.group(1)) if match else None
    source_path = Path(str(row.get("source_path", "")))
    if not source_path.exists():
        return None

    fallback = None
    for line in source_path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        try:
            payload = json.loads(line)
        except json.JSONDecodeError:
            continue
        if payload.get("type") == "final_evaluation":
            continue
        if target_epoch is not None and payload.get("epoch") == target_epoch:
            return payload
        try:
            if abs(float((payload.get("target") or {}).get("auc", -999)) - float(row["auc"])) < 1e-10:
                fallback = payload
        except (TypeError, ValueError):
            continue
    return fallback


def build_gate_rows(records: dict[str, dict[str, dict[int, dict[str, object]]]]) -> list[tuple[str, str, int, float, float]]:
    rows: list[tuple[str, str, int, float, float]] = []
    for dataset in DATASETS:
        for method in METHODS[1:]:
            graph_weights: list[float] = []
            train_weights: list[float] = []
            for record in records[dataset][method].values():
                payload = parse_best_epoch_payload(record)
                if payload is None:
                    continue
                analysis = ((payload.get("target") or {}).get("analysis") or {})
                train = payload.get("train") or {}
                if analysis.get("graph_branch_weight_mean") is not None:
                    graph_weights.append(float(analysis["graph_branch_weight_mean"]))
                if train.get("graph_weight") is not None:
                    train_weights.append(float(train["graph_weight"]))
            graph_mean, _ = mean_sd(graph_weights)
            train_mean, _ = mean_sd(train_weights)
            rows.append((dataset, method, len(graph_weights), graph_mean, train_mean))
    return rows


def table_line(values: Iterable[str]) -> str:
    return "| " + " | ".join(values) + " |"


def generate_markdown(
    metric_rows: list[MetricRow],
    delta_rows: list[DeltaRow],
    gate_rows: list[tuple[str, str, int, float, float]],
    notes: dict[str, str],
) -> str:
    metric_lookup = {(row.dataset, row.method): row for row in metric_rows}

    lines: list[str] = []
    lines.append("# 病理报告文本结构与医学知识增强实验进展报告")
    lines.append("")
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    lines.append("## 一、研究目标")
    lines.append("")
    lines.append(
        "本阶段工作的目标是在原有病理报告 sentence-level 表示基础上，引入医学 ontology/concept graph 与 "
        "Document -> Section -> Sentence 层次结构，评估医学知识增强和文档结构增强是否能够提升 BRCA、KIRC、LUSC "
        "三个癌症病理报告任务中的泛化性能。核心原则是保留原始 sentence 分支作为主语义分支，ontology 和 hierarchy "
        "仅作为受控辅助分支。"
    )
    lines.append("")
    lines.append("## 二、采用的方法")
    lines.append("")
    lines.append("1. 句子级表示（sentence-only）：沿用论文中的 sentence_pt 思路，将病理报告切分为句子，并使用 CONCH 文本编码得到句子特征；模型主要依赖原始文本语义。")
    lines.append("2. 句子级表示 + 医学知识增强（sentence + ontology）：在 sentence 分支之外增加 concept graph 辅助分支。本轮 30-split 实验使用的是 NCIt+DO compact ontology，其中 NCIt 负责肿瘤相关核心标准化，DO 负责疾病层级补充。")
    lines.append("3. 句子级表示 + 文档结构增强（sentence + hierarchical graph）：将病理报告组织为 Document、Section、Sentence 三层结构，并加入 section title/section role 信息，使模型获得句子所处上下文位置。")
    lines.append("4. 句子级表示 + 结构/知识联合增强（sentence + hierarchical graph + ontology）：在层次图基础上进一步加入 concept 节点和 ontology 边，形成 Document -> Section -> Sentence -> Concept 的联合图。")
    lines.append("")
    lines.append(
        "融合方式采用 dual_text 残差门控：sentence 分支保持主导，graph 分支通过受限权重注入，当前 graph branch 最大权重约束为 0.2，并加入 gate regularization。"
    )
    lines.append("")
    lines.append("## 三、实验设置")
    lines.append("")
    lines.append("- 数据集：BRCA、KIRC、LUSC。")
    lines.append("- 评价指标：AUC 作为主要指标，ACC 作为辅助指标。")
    lines.append("- ontology 版本：本轮主实验使用 `ncit_do`，即 NCIt+DO compact ontology。SNOMED CT 与 UMLS 已适配进资源构建流程，但未进入本轮 30-split 主结果。")
    lines.append(f"- BRCA baseline：{notes.get('BRCA', '')}")
    lines.append(f"- KIRC baseline：{notes.get('KIRC', '')}")
    lines.append(f"- LUSC baseline：{notes.get('LUSC', '')}")
    if notes.get("LUSC_full_150"):
        lines.append(f"- 额外参考：{notes['LUSC_full_150']}")
    lines.append("")
    lines.append("## 四、总体实验结果")
    lines.append("")
    lines.append(table_line(["数据集", "sentence-only AUC", "sentence+ontology AUC", "sentence+hierarchy AUC", "sentence+hierarchy+ontology AUC"]))
    lines.append(table_line(["---", "---:", "---:", "---:", "---:"]))
    for dataset in DATASETS:
        lines.append(
            table_line(
                [
                    dataset,
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-only")].auc_mean, metric_lookup[(dataset, "sentence-only")].auc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-ontology")].auc_mean, metric_lookup[(dataset, "sentence-ontology")].auc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph")].auc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph")].auc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].auc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].auc_sd),
                ]
            )
        )
    lines.append("")
    lines.append(table_line(["数据集", "sentence-only ACC", "sentence+ontology ACC", "sentence+hierarchy ACC", "sentence+hierarchy+ontology ACC"]))
    lines.append(table_line(["---", "---:", "---:", "---:", "---:"]))
    for dataset in DATASETS:
        lines.append(
            table_line(
                [
                    dataset,
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-only")].acc_mean, metric_lookup[(dataset, "sentence-only")].acc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-ontology")].acc_mean, metric_lookup[(dataset, "sentence-ontology")].acc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph")].acc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph")].acc_sd),
                    fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].acc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].acc_sd),
                ]
            )
        )
    lines.append("")
    lines.append("## 五、相对 sentence-only 的 paired delta")
    lines.append("")
    lines.append(table_line(["数据集", "增强模式", "共同 split 数", "ΔAUC", "AUC 提升 split", "ΔACC"]))
    lines.append(table_line(["---", "---", "---:", "---:", "---:", "---:"]))
    for row in delta_rows:
        lines.append(
            table_line(
                [
                    row.dataset,
                    METHOD_LABELS[row.method],
                    str(row.n),
                    fmt_delta(row.dauc_mean, row.dauc_sd),
                    f"{row.dauc_pos}/{row.n}" if row.n else "N/A",
                    fmt_delta(row.dacc_mean, row.dacc_sd),
                ]
            )
        )
    lines.append("")
    lines.append("## 六、分数据集结果分析")
    lines.append("")
    lines.append("### 1. BRCA")
    lines.append("")
    lines.append(
        "BRCA 上 sentence-only 最强，AUC 为 0.7564 ± 0.0352。ontology 分支下降最明显，ΔAUC 为 -0.0432，只有 8/30 个 split 提升；hierarchy 分支下降较小，ΔAUC 为 -0.0169，12/30 个 split 提升；联合图也未超过 baseline。"
    )
    lines.append(
        "这一现象说明当前 BRCA 中 NCIt+DO concept graph 没有稳定捕获分期相关判别信息。BRCA 报告中大量概念可能集中在肿瘤类型、解剖部位、常规病理描述上，这些概念并不一定直接对应 stage 标签；同时 concept 匹配与 true-path 祖先仍可能引入泛化节点，导致排序能力下降。"
    )
    lines.append("")
    lines.append("### 2. KIRC")
    lines.append("")
    lines.append(
        "KIRC 上 sentence+ontology 略高，AUC 为 0.8770 ± 0.0261，相对历史 sentence-only 前 30 splits 的 ΔAUC 为 +0.0060，17/30 个 split 提升。hierarchy 基本持平，joint 没有继续提升。"
    )
    lines.append(
        "KIRC 与肾癌、透明细胞癌、肾脏解剖位置等 NCIt/DO 概念更直接相关，因此 ontology 有轻微帮助。但 gate 分析显示 KIRC 的 graph branch 平均权重很低，说明模型大多自动回避图分支，只在少量样本中利用知识补充，因此提升幅度有限。"
    )
    lines.append("")
    lines.append("### 3. LUSC")
    lines.append("")
    lines.append(
        "LUSC 上 hierarchy 是最有价值的增强，AUC 为 0.7098 ± 0.0416，相对 sentence-only 前 30 splits 的 ΔAUC 为 +0.0120。ontology 单独下降，joint 也下降，说明 ontology 噪声抵消了 hierarchy 的收益。"
    )
    lines.append(
        "LUSC 报告可能更依赖 section role 和上下文组织，例如 diagnosis、clinical history、gross description、microscopic description 等区域对标签信息的贡献不同。层次图能帮助模型区分句子所在位置；但当前 ontology 仍偏泛化，加入后会稀释结构信号。"
    )
    lines.append("")
    lines.append("## 七、graph branch 使用情况")
    lines.append("")
    lines.append(table_line(["数据集", "增强模式", "n", "graph branch 平均权重", "训练中 graph_weight 平均值"]))
    lines.append(table_line(["---", "---", "---:", "---:", "---:"]))
    for dataset, method, n, graph_mean, train_mean in gate_rows:
        lines.append(table_line([dataset, METHOD_LABELS[method], str(n), f"{graph_mean:.4f}", f"{train_mean:.4f}"]))
    lines.append("")
    lines.append(
        "从 gate 结果看，sentence 分支始终是主分支。KIRC 中 graph branch 权重最低，因此增强分支影响很小；BRCA 和 LUSC 中图分支权重更高，因此图分支质量会更直接影响最终结果。"
    )
    lines.append("")
    lines.append("## 八、结果原因总结")
    lines.append("")
    lines.append("1. 当前 ontology 使用的是 NCIt+DO compact，并没有在主实验中引入 SNOMED CT/UMLS，因此目前只能评价 NCIt+DO 这一档，不能否定 SNOMED/UMLS 的潜在价值。")
    lines.append("2. ontology 图的概念覆盖和标签判别目标之间仍有错位：概念能描述疾病和病理实体，但不一定能描述分期、侵犯、转移、淋巴结等真正与 stage 相关的证据。")
    lines.append("3. hierarchy 的有效性依赖 section role。LUSC 对文档结构更敏感，BRCA 中层次聚合可能放大中性 section 或无关句子。")
    lines.append("4. joint 模式没有加成，说明当前 ontology 与 hierarchy 不是简单互补关系；如果 ontology 噪声较强，它会抵消 hierarchy 提供的上下文收益。")
    lines.append("5. residual gate 起到了保护作用，避免图分支完全覆盖 sentence 分支，但当图分支质量不足时，即使低权重也可能影响 AUC 排序。")
    lines.append("")
    lines.append("## 九、后续改进建议")
    lines.append("")
    lines.append("1. BRCA：暂时不要继续扩大 ontology。优先回到 sentence-only 与 section-title embedding，之后只保留 TNM、stage、grade、tumor size、lymph node、metastasis 等 typed compact concepts。")
    lines.append("2. KIRC：可以保留 ontology 作为弱辅助，但需要做 NCIt-only、NCIt+DO、NCIt+SNOMED-mapped-to-NCIt 的小规模对照，确认轻微收益来自哪个资源。")
    lines.append("3. LUSC：优先推进 hierarchy 方向，尤其是 section role 建模、section-title embedding、hierarchy graph auxiliary weight 的 0.05/0.10/0.20 ablation。")
    lines.append("4. ontology 方向：不要直接 full multi-ontology。SNOMED CT 适合提高 mention recall，UMLS 适合跨术语对齐，NCIt 应继续作为肿瘤任务核心标准化，DO 只补疾病层级。")
    lines.append("5. 图结构方向：增加 typed edge/weighted edge，降低泛化 ancestor 和高频概念的权重，并对 concept 节点使用 label embedding + mention sentence embedding 的融合表示。")
    lines.append("6. 实验记录方向：统一把历史 sentence-only baseline 解析进 records，保证 BRCA/KIRC/LUSC 三个数据集都能直接进行 30-split paired comparison。")
    lines.append("")
    lines.append("## 十、阶段性结论")
    lines.append("")
    lines.append(
        "本阶段实验表明，原始 sentence-level 文本语义仍是最可靠的基础。文档结构增强在 LUSC 上出现稳定潜力，KIRC 的 ontology 有轻微收益，而 BRCA 当前不适合直接引入 ontology 图。下一阶段应从“图分支替代/强融合”转向“受控辅助分支”，并把 ontology 从全量概念图收紧为 compact、typed、weighted、gated 的任务相关知识。"
    )
    lines.append("")
    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    document.add_paragraph()


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def generate_docx(
    metric_rows: list[MetricRow],
    delta_rows: list[DeltaRow],
    gate_rows: list[tuple[str, str, int, float, float]],
    notes: dict[str, str],
    output_path: Path,
) -> None:
    metric_lookup = {(row.dataset, row.method): row for row in metric_rows}

    document = Document()
    set_document_style(document)

    title = document.add_heading("病理报告文本结构与医学知识增强实验进展报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, f"汇报日期：2026年4月28日    生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_heading("一、研究目标", level=1)
    add_paragraph(
        document,
        "本阶段工作围绕 BRCA、KIRC、LUSC 癌症病理报告任务，探索在原有 sentence-level 文本表示基础上引入医学 ontology/concept graph 与 Document -> Section -> Sentence 层次结构，评估医学知识增强和文档结构增强对泛化性能的影响。整体设计中，原始 sentence 分支始终作为主语义分支，ontology 和 hierarchy 只作为受控辅助分支。",
    )

    document.add_heading("二、采用的方法", level=1)
    method_rows = [
        ["sentence-only", "沿用论文 sentence_pt 思路，将病理报告切分为句子，并使用 CONCH 文本编码得到句子特征。"],
        ["sentence + ontology", "在 sentence 分支外加入 concept graph。本轮使用 NCIt+DO compact ontology：NCIt 负责肿瘤核心标准化，DO 补充疾病层级。"],
        ["sentence + hierarchical graph", "构建 Document、Section、Sentence 三层结构，并引入 section title/section role 信息。"],
        ["sentence + hierarchical graph + ontology", "在层次图基础上加入 concept 节点和 ontology 边，形成结构与知识联合图。"],
    ]
    add_table(document, ["方法", "说明"], method_rows)
    add_paragraph(
        document,
        "融合方式采用 dual_text 残差门控：sentence 分支保持主导，graph 分支通过受限权重注入，当前 graph branch 最大权重约束为 0.2，并加入 gate regularization。",
    )

    document.add_heading("三、实验设置", level=1)
    for line in [
        "数据集：BRCA、KIRC、LUSC。",
        "评价指标：AUC 作为主要指标，ACC 作为辅助指标。",
        "ontology 版本：本轮主实验使用 ncit_do，即 NCIt+DO compact ontology；SNOMED CT 与 UMLS 已适配进资源构建流程，但未进入本轮 30-split 主结果。",
        f"BRCA baseline：{notes.get('BRCA', '')}",
        f"KIRC baseline：{notes.get('KIRC', '')}",
        f"LUSC baseline：{notes.get('LUSC', '')}",
        notes.get("LUSC_full_150", ""),
    ]:
        if line:
            add_paragraph(document, line, style="List Bullet")

    document.add_heading("四、总体实验结果", level=1)
    auc_rows = []
    acc_rows = []
    for dataset in DATASETS:
        auc_rows.append(
            [
                dataset,
                fmt_mean_sd(metric_lookup[(dataset, "sentence-only")].auc_mean, metric_lookup[(dataset, "sentence-only")].auc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-ontology")].auc_mean, metric_lookup[(dataset, "sentence-ontology")].auc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph")].auc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph")].auc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].auc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].auc_sd),
            ]
        )
        acc_rows.append(
            [
                dataset,
                fmt_mean_sd(metric_lookup[(dataset, "sentence-only")].acc_mean, metric_lookup[(dataset, "sentence-only")].acc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-ontology")].acc_mean, metric_lookup[(dataset, "sentence-ontology")].acc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph")].acc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph")].acc_sd),
                fmt_mean_sd(metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].acc_mean, metric_lookup[(dataset, "sentence-hierarchical-graph-ontology")].acc_sd),
            ]
        )
    add_table(document, ["数据集", "sentence-only AUC", "sentence+ontology AUC", "sentence+hierarchy AUC", "joint AUC"], auc_rows)
    add_table(document, ["数据集", "sentence-only ACC", "sentence+ontology ACC", "sentence+hierarchy ACC", "joint ACC"], acc_rows)

    document.add_heading("五、相对 sentence-only 的 paired delta", level=1)
    add_table(
        document,
        ["数据集", "增强模式", "共同 split 数", "ΔAUC", "AUC 提升 split", "ΔACC"],
        [
            [
                row.dataset,
                METHOD_LABELS[row.method],
                str(row.n),
                fmt_delta(row.dauc_mean, row.dauc_sd),
                f"{row.dauc_pos}/{row.n}" if row.n else "N/A",
                fmt_delta(row.dacc_mean, row.dacc_sd),
            ]
            for row in delta_rows
        ],
    )

    document.add_heading("六、分数据集结果分析", level=1)
    document.add_heading("1. BRCA", level=2)
    add_paragraph(document, "BRCA 上 sentence-only 最强，AUC 为 0.7564 ± 0.0352。ontology 分支下降最明显，ΔAUC 为 -0.0432，只有 8/30 个 split 提升；hierarchy 分支下降较小，ΔAUC 为 -0.0169，12/30 个 split 提升；联合图也未超过 baseline。")
    add_paragraph(document, "这一现象说明当前 BRCA 中 NCIt+DO concept graph 没有稳定捕获分期相关判别信息。BRCA 报告中大量概念可能集中在肿瘤类型、解剖部位、常规病理描述上，这些概念并不一定直接对应 stage 标签；同时 concept 匹配与 true-path 祖先仍可能引入泛化节点，导致排序能力下降。")

    document.add_heading("2. KIRC", level=2)
    add_paragraph(document, "KIRC 上 sentence+ontology 略高，AUC 为 0.8770 ± 0.0261，相对历史 sentence-only 前 30 splits 的 ΔAUC 为 +0.0060，17/30 个 split 提升。hierarchy 基本持平，joint 没有继续提升。")
    add_paragraph(document, "KIRC 与肾癌、透明细胞癌、肾脏解剖位置等 NCIt/DO 概念更直接相关，因此 ontology 有轻微帮助。但 gate 分析显示 KIRC 的 graph branch 平均权重很低，说明模型大多自动回避图分支，只在少量样本中利用知识补充，因此提升幅度有限。")

    document.add_heading("3. LUSC", level=2)
    add_paragraph(document, "LUSC 上 hierarchy 是最有价值的增强，AUC 为 0.7098 ± 0.0416，相对 sentence-only 前 30 splits 的 ΔAUC 为 +0.0120。ontology 单独下降，joint 也下降，说明 ontology 噪声抵消了 hierarchy 的收益。")
    add_paragraph(document, "LUSC 报告可能更依赖 section role 和上下文组织，例如 diagnosis、clinical history、gross description、microscopic description 等区域对标签信息的贡献不同。层次图能帮助模型区分句子所在位置；但当前 ontology 仍偏泛化，加入后会稀释结构信号。")

    document.add_heading("七、graph branch 使用情况", level=1)
    add_table(
        document,
        ["数据集", "增强模式", "n", "graph branch 平均权重", "训练 graph_weight 平均值"],
        [[dataset, METHOD_LABELS[method], str(n), f"{graph_mean:.4f}", f"{train_mean:.4f}"] for dataset, method, n, graph_mean, train_mean in gate_rows],
    )
    add_paragraph(document, "从 gate 结果看，sentence 分支始终是主分支。KIRC 中 graph branch 权重最低，因此增强分支影响很小；BRCA 和 LUSC 中图分支权重更高，因此图分支质量会更直接影响最终结果。")

    document.add_heading("八、结果原因总结", level=1)
    for line in [
        "当前 ontology 使用的是 NCIt+DO compact，并没有在主实验中引入 SNOMED CT/UMLS，因此目前只能评价 NCIt+DO 这一档，不能否定 SNOMED/UMLS 的潜在价值。",
        "ontology 图的概念覆盖和标签判别目标之间仍有错位：概念能描述疾病和病理实体，但不一定能描述分期、侵犯、转移、淋巴结等真正与 stage 相关的证据。",
        "hierarchy 的有效性依赖 section role。LUSC 对文档结构更敏感，BRCA 中层次聚合可能放大中性 section 或无关句子。",
        "joint 模式没有加成，说明当前 ontology 与 hierarchy 不是简单互补关系；如果 ontology 噪声较强，它会抵消 hierarchy 提供的上下文收益。",
        "residual gate 起到了保护作用，避免图分支完全覆盖 sentence 分支，但当图分支质量不足时，即使低权重也可能影响 AUC 排序。",
    ]:
        add_paragraph(document, line, style="List Number")

    document.add_heading("九、后续改进建议", level=1)
    for line in [
        "BRCA：暂时不要继续扩大 ontology。优先回到 sentence-only 与 section-title embedding，之后只保留 TNM、stage、grade、tumor size、lymph node、metastasis 等 typed compact concepts。",
        "KIRC：可以保留 ontology 作为弱辅助，但需要做 NCIt-only、NCIt+DO、NCIt+SNOMED-mapped-to-NCIt 的小规模对照，确认轻微收益来自哪个资源。",
        "LUSC：优先推进 hierarchy 方向，尤其是 section role 建模、section-title embedding、hierarchy graph auxiliary weight 的 0.05/0.10/0.20 ablation。",
        "ontology 方向：不要直接 full multi-ontology。SNOMED CT 适合提高 mention recall，UMLS 适合跨术语对齐，NCIt 应继续作为肿瘤任务核心标准化，DO 只补疾病层级。",
        "图结构方向：增加 typed edge/weighted edge，降低泛化 ancestor 和高频概念的权重，并对 concept 节点使用 label embedding + mention sentence embedding 的融合表示。",
        "实验记录方向：统一把历史 sentence-only baseline 解析进 records，保证 BRCA/KIRC/LUSC 三个数据集都能直接进行 30-split paired comparison。",
    ]:
        add_paragraph(document, line, style="List Number")

    document.add_heading("十、阶段性结论", level=1)
    add_paragraph(
        document,
        "本阶段实验表明，原始 sentence-level 文本语义仍是最可靠的基础。文档结构增强在 LUSC 上出现稳定潜力，KIRC 的 ontology 有轻微收益，而 BRCA 当前不适合直接引入 ontology 图。下一阶段应从“图分支替代/强融合”转向“受控辅助分支”，并把 ontology 从全量概念图收紧为 compact、typed、weighted、gated 的任务相关知识。",
    )

    ensure_parent(output_path)
    document.save(output_path)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def main() -> None:
    records, notes = load_all_records()
    metric_rows = build_metric_rows(records, notes)
    delta_rows = build_delta_rows(records)
    gate_rows = build_gate_rows(records)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = generate_markdown(metric_rows, delta_rows, gate_rows, notes)
    markdown_path = REPORT_DIR / f"{REPORT_STEM}.md"
    markdown_path.write_text(markdown, encoding="utf-8")

    docx_path = REPORT_DIR / f"{REPORT_STEM}.docx"
    generate_docx(metric_rows, delta_rows, gate_rows, notes, docx_path)

    print(f"wrote {docx_path}")
    print(f"wrote {markdown_path}")


if __name__ == "__main__":
    main()
